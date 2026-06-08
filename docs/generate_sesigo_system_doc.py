"""
Generator for the Sesigo Data Portal - System Explanation Document (Word).

This script builds docs/Sesigo_Data_Portal_System_Explanation_Document.docx
directly from the verified BONASOV1 implementation (backend Django apps +
Next.js frontend routes). Re-run after schema/feature changes to keep the
document aligned with the real system.

Usage:
    backend/venv/bin/python docs/generate_sesigo_system_doc.py
"""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

OUT = Path(__file__).resolve().parent / "Sesigo_Data_Portal_System_Explanation_Document.docx"

BLUE_DARK = RGBColor(0x1F, 0x4E, 0x79)
BLUE = RGBColor(0x2E, 0x75, 0xB6)
GREY = RGBColor(0x80, 0x80, 0x80)
HDR_FILL = "D9E2F3"


def set_cell_bg(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_footer(document):
    section = document.sections[-1]
    footer = section.footer
    p = footer.paragraphs[0]
    p.text = "Powered by BONASO | Prepared as senior developer documentation"
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = GREY


def add_header(document):
    section = document.sections[-1]
    header = section.header
    p = header.paragraphs[0]
    p.text = "Sesigo Data Portal - System Explanation Document"
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in p.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = GREY


def h1(document, text):
    p = document.add_heading(level=1)
    run = p.add_run(text)
    run.font.color.rgb = BLUE_DARK
    run.font.size = Pt(18)
    run.bold = True
    return p


def h2(document, text):
    p = document.add_heading(level=2)
    run = p.add_run(text)
    run.font.color.rgb = BLUE
    run.font.size = Pt(13)
    run.bold = True
    return p


def para(document, text, size=10.5):
    p = document.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    return p


def bullets(document, items):
    for it in items:
        p = document.add_paragraph(style="List Bullet")
        run = p.add_run(it)
        run.font.size = Pt(10.5)


def numbered(document, items):
    for it in items:
        p = document.add_paragraph(style="List Number")
        run = p.add_run(it)
        run.font.size = Pt(10.5)


def table(document, headers, rows, widths=None):
    t = document.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(htext)
        run.bold = True
        run.font.size = Pt(10)
        set_cell_bg(hdr[i], HDR_FILL)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9.5)
    if widths:
        for i, w in enumerate(widths):
            for r in t.rows:
                r.cells[i].width = Inches(w)
    document.add_paragraph()
    return t


def diagram_flow(document, steps):
    """Render a simple vertical flow as centered text with arrows (no image deps)."""
    for i, step in enumerate(steps):
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(step)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = BLUE_DARK
        if i < len(steps) - 1:
            a = document.add_paragraph()
            a.alignment = WD_ALIGN_PARAGRAPH.CENTER
            ar = a.add_run("↓")
            ar.font.size = Pt(11)
    document.add_paragraph()


def caption(document, text):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = GREY


def build():
    doc = Document()

    # Base style
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    add_header(doc)
    add_footer(doc)

    # ---- Title block ----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Sesigo Data Portal")
    r.bold = True
    r.font.size = Pt(30)
    r.font.color.rgb = BLUE_DARK

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("System Explanation and Component Map")
    r.bold = True
    r.font.size = Pt(17)
    r.font.color.rgb = BLUE

    tag = doc.add_paragraph()
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tag.add_run("What it is, how it works, and how it is used")
    r.italic = True
    r.font.size = Pt(12)

    para(doc,
         "This edition is generated directly from the current BONASOV1 implementation "
         "(Django REST backend apps and the Next.js App Router frontend) so that the "
         "explanation matches the system as actually built.")

    # ---- Document Control ----
    h1(doc, "Document Control")
    table(doc, ["Item", "Description"], [
        ["Document purpose", "Explain the Sesigo Data Portal in business, user, data, and technical terms."],
        ["Audience", "BONASO management, system developers, M&E managers and officers, coordinators, sub-grantees, funders/clients, and system users."],
        ["Scope", "Portal purpose, programme hierarchy, user roles, full feature set, data workflow, data-quality and review controls, live/training separation, messaging/collaboration, imports/exports, core components, and technical architecture."],
        ["Source basis", "Verified from the BONASOV1 codebase: backend Django apps (users, organizations, indicators, projects, respondents, aggregates, events, social, flags, analysis, profiles, uploads, messaging) and the Next.js frontend route map."],
        ["Naming standard", "Sesigo Data Portal for the platform, Sesigo Live System for production, Sesigo Training Mode for training, and Powered by BONASO as the footer/tagline."],
    ], widths=[1.8, 4.7])

    h2(doc, "Contents")
    contents = [
        "1. Executive summary",
        "2. What the Sesigo Data Portal is",
        "3. Why the portal exists",
        "4. Who uses the portal",
        "5. Programme and funding structure",
        "6. Main system components",
        "7. How the portal works from setup to dashboard",
        "8. Technical architecture",
        "9. User roles and permissions",
        "10. Data capture, indicators, and disaggregation",
        "11. Data quality, review, and flagging",
        "12. Reports, dashboards, analytics, and exports",
        "13. Imports, exports, and report workbooks",
        "14. Messaging, announcements, notifications, and social tracking",
        "15. Sesigo Live System vs Sesigo Training Mode",
        "16. Simplified database model",
        "17. How the system is used by each user group",
        "18. Administration, security, and quality assurance",
        "19. Implementation checklist for developers",
        "20. Future improvements",
        "21. Architectural conclusion",
    ]
    for c in contents:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(c).font.size = Pt(10.5)

    doc.add_page_break()

    # ---- 1. Executive Summary ----
    h1(doc, "1. Executive Summary")
    para(doc,
         "The Sesigo Data Portal is a project-scoped information management and reporting platform used by "
         "BONASO to manage funded civil society organisation programme data. It brings together clients/funders, "
         "projects, BONASO oversight, coordinator organisations, sub-grantees, assigned indicators, data capture, "
         "data-quality review, dashboards, exports, and analytics in one controlled system.")
    para(doc,
         "The portal is not only a data-entry tool. It is a structured reporting environment that controls who can "
         "report, what they can report, which project they report under, and whether the data belongs to the official "
         "live system or the training environment. Around that core it adds collaboration and governance features: "
         "data-quality flags, internal messaging and announcements, notifications, custom respondent profiles, "
         "assessments, social-media tracking, bulk imports, generated exports, scheduled reports, and a user-activity "
         "audit trail.")
    para(doc,
         "The core design rule is simple: a client funds a project, BONASO oversees the project, coordinator "
         "organisations manage delivery, sub-grantees capture assigned indicators, submitted data is reviewed and "
         "approved, and approved data feeds dashboards, reports, exports, and analytics.")
    para(doc,
         "A second core rule is that training data must never mix with live data. Sesigo Training Mode uses /training "
         "routes and a mode-aware request signal (training_only=true / mode=training). Sesigo Live System uses normal "
         "production routes. The backend remains the final authority through apply_training_filter(...) on reads and "
         "assert_project_write_allowed(...) on writes, with Project.is_training as the source of truth.")

    # ---- 2. What it is ----
    h1(doc, "2. What the Sesigo Data Portal Is")
    para(doc,
         "The Sesigo Data Portal is a web-based data management system for planning, capturing, reviewing, and "
         "analysing project implementation data across BONASO-supported programmes. It provides a single place where "
         "project structures, organisation roles, reporting assignments, indicator results, respondents and events, "
         "narrative reports, dashboards, and exports are managed.")
    para(doc,
         "From a business point of view, Sesigo is a programme coordination and reporting platform. From a technical "
         "point of view, it is a Next.js App Router frontend connected to a Django REST Framework API and a relational "
         "database, with JWT authentication, role and organisation scoping, controlled file uploads, and reporting "
         "outputs.")
    h2(doc, "2.1 The portal in one sentence")
    para(doc,
         "Sesigo Data Portal helps BONASO and its partners manage funded project reporting by connecting projects, "
         "organisations, indicators, data capture, validation, review, and dashboards under one governed digital "
         "system.")

    # ---- 3. Why ----
    h1(doc, "3. Why the Portal Exists")
    para(doc,
         "The portal exists to solve fragmented data management. Without a central system, data is captured in "
         "different Excel files, emails, paper tools, or organisation-specific templates. This creates inconsistent "
         "indicator names, duplicate records, weak version control, delayed reporting, and difficulty seeing progress "
         "across coordinators and sub-grantees.")
    para(doc,
         "Sesigo improves this by standardising programme structure, assigning indicators to the correct "
         "organisations, enforcing project scope, supporting review and flagging before data is approved, and producing "
         "dashboards and exports from approved operational data.")
    h2(doc, "3.1 Problems addressed")
    bullets(doc, [
        "Fragmented reporting across multiple organisations and templates.",
        "Difficulty linking data to the correct project, funder, organisation, and indicator.",
        "Indicator duplication caused by different wording for similar measures (addressed by canonical indicators plus indicator aliases).",
        "Weak visibility of coordinator and sub-grantee performance.",
        "Risk of training or demo records contaminating official reporting data.",
        "Slow reporting caused by manual consolidation and repeated Excel clean-up.",
        "No shared trail of who created, edited, approved, exported, or imported records.",
    ])
    h2(doc, "3.2 Value delivered")
    bullets(doc, [
        "Centralised information management for BONASO and funded partners.",
        "Clear project and organisation accountability.",
        "Better data quality through validation, flagging, review, and approval workflows.",
        "Real-time or near real-time dashboards for management and funder reporting.",
        "Consistent reporting structure across live operations and training practice.",
        "An audit trail (user activity) and in-app communication that keep teams aligned.",
    ])

    # ---- 4. Who uses ----
    h1(doc, "4. Who Uses the Portal")
    para(doc,
         "The portal is used by different people at different levels of responsibility. A single user has one "
         "platform-level role (users.User.role) and participates in projects through one or more project-scoped "
         "organisation roles (projects.ProjectOrganization.role).")
    table(doc, ["User group", "Main purpose in the portal", "Typical activities"], [
        ["BONASO admin / system administrator", "Configure and maintain the platform.",
         "Create users, manage roles and permissions, configure projects, assign organisations and hierarchy, manage indicators and targets, monitor data quality and flags, run training reset/cleanup, review activity logs, manage settings."],
        ["M&E Manager (manager)", "Oversee project delivery and reporting.",
         "Review and approve submissions, set coordinator targets, monitor indicators, resolve flags, follow up coordinators, view dashboards, schedule and export reports."],
        ["M&E Officer (officer)", "Operational data work in allowed scope.",
         "Capture and edit operational records, raise and action flags, prepare reports, manage events and respondents within assigned scope."],
        ["Coordinator organisation", "Coordinate delivery under a project.",
         "Manage sub-grantees, report own assigned indicators, review sub-grantee progress, monitor flags, submit narrative reports."],
        ["Sub-grantee organisation", "Capture project implementation data.",
         "Enter aggregate data, capture respondents/interactions and event participation where applicable, upload narrative reports, view own reporting status and flags."],
        ["Data Collector (collector)", "Capture field-level data in allowed scope.",
         "Submit data against assigned indicators and projects; use batch capture and check-in."],
        ["Client / Funder (client)", "View approved project performance.",
         "Access read-only reports, dashboards, exports, and high-level progress summaries."],
    ], widths=[1.9, 1.9, 2.7])

    # ---- 5. Programme structure ----
    h1(doc, "5. Programme and Funding Structure")
    para(doc,
         "The Sesigo programme model starts with a client/funder. The client funds a project. The project is overseen "
         "by BONASO. Implementation happens through coordinator organisations, which may manage sub-grantees. A "
         "coordinator can also implement some indicators directly when the project assigns it that responsibility.")
    diagram_flow(doc, [
        "Client / Funder (ClientOrganization)  — funds →",
        "Project (reporting boundary, is_training flag)  — overseen by →",
        "BONASO Overseer (governance + accountability)  — governs / assigns delivery →",
        "Coordinator Organisations (manage sub-grantees, may self-report)",
        "Sub-grantees + Coordinator direct implementation  →  Assigned indicators + targets",
    ])
    caption(doc, "Figure 1: Client, project, BONASO, coordinator, and sub-grantee hierarchy.")
    h2(doc, "5.1 Explanation of the structure")
    table(doc, ["Component", "Meaning in Sesigo", "Why it matters"], [
        ["Client / Funder (ClientOrganization)", "The organisation or funding source that finances a project.",
         "Links project performance and reporting back to the funding source."],
        ["Project", "The official reporting boundary for a funded programme; holds status and the is_training flag.",
         "All targets, indicators, organisations, dashboards, and training/live status are scoped through the project."],
        ["BONASO Overseer", "The organisation responsible for governance, oversight, coordination, and accountability.",
         "Ensures project delivery is monitored and reported consistently."],
        ["Coordinator organisation", "A lead implementing organisation under the project.",
         "Coordinates sub-grantees, may submit own reports, and helps manage delivery."],
        ["Sub-grantee organisation", "An organisation managed by a coordinator and assigned specific indicators.",
         "Captures service-delivery data and contributes to project results."],
        ["Coordinator direct implementation", "A coordinator reporting its own assigned indicators (is_implementer / can_report_indicators).",
         "Supports cases where the coordinator is both manager and implementer."],
        ["ProjectOrganizationHierarchy", "Explicit parent-child links between organisations inside a project.",
         "Defines who a coordinator can manage and view; backed by Project.hierarchy_overrides as a fallback."],
    ], widths=[1.9, 2.5, 2.1])

    # ---- 6. Components ----
    h1(doc, "6. Main System Components")
    para(doc,
         "The portal combines business components (programme realities such as projects, organisations, indicators, "
         "targets, respondents, events, reports, and dashboards) with technical components that ensure secure access, "
         "validation, storage, training isolation, communication, and exports. The list below reflects the actual "
         "Django apps in the backend.")
    table(doc, ["Component (app/model)", "What it does", "Important rules"], [
        ["Clients / Funders (ClientOrganization)", "Stores the funder/client linked to projects.", "A project should be traceable to its funding source."],
        ["Projects (Project)", "Defines the reporting boundary, status, and training/live flag.", "Every major record links to a project; is_training is the mode source of truth."],
        ["Organisations (Organization)", "Stores BONASO, coordinators, sub-grantees, and implementing partners.", "Users and reports are scoped through organisation relationships and the org tree."],
        ["Project organisations (ProjectOrganization)", "Defines how an organisation participates in a specific project (role, is_coordinator, is_sub_grantee, is_implementer, can_report_indicators).", "The same organisation can have different roles in different projects."],
        ["Organisation hierarchy (ProjectOrganizationHierarchy)", "Parent-child relationships between coordinators and sub-grantees.", "A coordinator should only manage/view allowed sub-grantee data."],
        ["Indicators (Indicator, IndicatorAlias)", "Standard performance measures plus alternate import/display names that resolve to a canonical indicator.", "Indicator wording is controlled; aliases reduce duplicates on import."],
        ["Assessments (Assessment, AssessmentIndicator)", "Group related indicators into a single data-collection instrument.", "Keeps multi-indicator collection consistent."],
        ["Project indicators (ProjectIndicator + targets/rules)", "Links indicators to a project with targets, per-organisation targets, and disaggregation rules.", "Targets are project-specific and time-bound where needed."],
        ["Indicator assignments (ProjectIndicatorAssignment)", "Controls which organisation may report which project indicator.", "Organisations should not report indicators they were not assigned."],
        ["Aggregates (Aggregate)", "Summary indicator values by period, project, organisation, and status.", "Records move through draft → pending → reviewed/flagged → approved/rejected before official dashboards."],
        ["Respondents (Respondent, Interaction, Response)", "Individual-level beneficiaries, their service interactions, and indicator responses.", "Handled with care; kept linked to respondent, project, event, and indicator."],
        ["Profiles (Profile, ProfileField)", "Extended/custom profile fields for respondents and participants.", "Allows configurable demographics and classification fields."],
        ["Events (Event, Participant, EventPhase)", "Activities, sessions, phases, and attendees.", "Event data follows the same project and training/live boundary."],
        ["Social (SocialPost)", "Social-media posts tracked against an indicator.", "Counts media/communication outputs within project scope."],
        ["Flags (Flag, FlagComment)", "Data-quality / follow-up / review flags attached to any record.", "Supports issue tracking, assignment, and resolution before approval."],
        ["Narrative reports (NarrativeReport)", "Supporting narrative reports and documents.", "Uploads are linked to project and organisation scope."],
        ["Analysis (Report, SavedQuery, ScheduledReport, CoordinatorTarget)", "Saved/generated reports, reusable queries, scheduled reports, and coordinator targets.", "Reads approved, scoped, correctly filtered data."],
        ["Uploads (Upload, ImportJob, ExportJob)", "File uploads plus tracked import and export jobs.", "Imports validate scope; exports respect role, organisation scope, and mode."],
        ["Messaging (Message, Announcement, Notification)", "Internal messages, scoped announcements, and system notifications.", "Keeps teams aligned without leaving the portal."],
        ["Users & audit (User, UserActivity)", "Accounts, platform roles, and a user-activity audit trail (login, create, update, delete, view, export, import).", "Backend permissions enforced for every read and write; actions are logged."],
    ], widths=[2.1, 2.4, 2.0])

    # ---- 7. Workflow ----
    h1(doc, "7. How the Portal Works from Setup to Dashboard")
    para(doc,
         "The system follows a controlled workflow. First the project structure is configured, then indicators and "
         "targets are assigned. Organisations capture data only against allowed indicators. The backend validates the "
         "submission. Records become pending, are reviewed or flagged, and approved data is used for reporting outputs.")
    diagram_flow(doc, [
        "1. Project setup — client, project, orgs, hierarchy, live/training",
        "2. Indicator setup — indicators, disaggregation, targets",
        "3. Assign reporting scope — project indicator assignment",
        "4. Capture data — aggregates, respondents, events, social, imports",
        "5. Validate — completeness, access, mode, assignment",
        "6. Review + flag + approve — draft/pending → reviewed/flagged → approved",
        "7. Report — dashboards, exports, analytics, narratives",
    ])
    caption(doc, "Figure 2: Data workflow from project setup to dashboards and exports.")
    h2(doc, "7.1 End-to-end process")
    table(doc, ["Step", "Process", "Description"], [
        ["1", "Project setup", "Create the funded project, define client/funder, status, dates, and whether it is live or training."],
        ["2", "Organisation setup", "Add BONASO oversight, coordinator organisations, sub-grantees, and their project-specific roles and hierarchy."],
        ["3", "Indicator setup", "Define indicators (and aliases/assessments), disaggregation rules, targets, and project-specific reporting expectations."],
        ["4", "Assignment", "Assign indicators to organisations so each coordinator or sub-grantee knows what to report."],
        ["5", "Data capture", "Users enter aggregates, respondent services/interactions, event participation, social posts, workbook/batch imports, or narrative reports."],
        ["6", "Validation", "System checks project scope, user scope, organisation scope, indicator assignment, and training/live mode."],
        ["7", "Review, flag, approve", "Authorised users review pending records, raise/resolve flags, and approve valid data (or reject)."],
        ["8", "Reporting", "Approved records feed dashboards, reports, analytics, scheduled reports, and exports."],
    ], widths=[0.6, 1.7, 4.2])

    # ---- 8. Technical architecture ----
    h1(doc, "8. Technical Architecture")
    para(doc,
         "The technical architecture separates the user interface, API logic, data storage, and reporting outputs. The "
         "Next.js frontend provides the user experience and a route-aware API client. The Django REST API enforces "
         "business rules, JWT authentication (SimpleJWT), role/organisation scope, and the live/training boundary. The "
         "relational database stores operational data. Media storage holds uploads and narrative reports. Dashboards "
         "and exports read from scoped, validated data.")
    diagram_flow(doc, [
        "Portal Users (admins, managers, officers, coordinators, sub-grantees, collectors, clients)",
        "Next.js Frontend — App Router + live/training routes",
        "Route-aware API client — adds training_only=true for /training",
        "Django REST API — JWT auth, business rules, write guards",
        "Access control — user role, org tree, project scope, mode filter",
        "Domain services — projects, orgs, indicators, aggregates, respondents, events",
        "Relational DB (operational data)  +  Media storage (uploads, narratives)",
        "Dashboards + Analytics + Exports — decision support",
    ])
    caption(doc, "Figure 3: High-level technical architecture.")
    h2(doc, "8.1 Technical component explanation")
    table(doc, ["Technical layer", "Description", "Responsibility"], [
        ["Next.js frontend", "App Router web interface with live and /training route trees.", "Renders forms, dashboards, reports, batch capture, and training routes."],
        ["Route-aware API client", "Frontend request handler aware of live vs training context.", "Adds training_only=true / mode=training when under /training routes."],
        ["Django REST API", "Backend service layer (DRF) exposing data and operations.", "Authenticates (JWT/Djoser), validates users, roles, projects, org scope, indicators, and live/training boundaries."],
        ["Access control layer", "Role, permissions, organisation tree, project scope, and mode filtering (organizations/access.py, projects training helpers).", "Prevents users from viewing or editing data outside their allowed scope."],
        ["Database", "Primary operational database (PostgreSQL in production; SQLite for local/dev).", "Stores users, organisations, projects, indicators, targets, aggregates, respondents, events, messages, flags, and audit records."],
        ["Uploads / media storage", "Storage for uploaded reports and supporting files; WhiteNoise serves static assets.", "Keeps narrative reports and attachments linked to projects and organisations."],
        ["Dashboards / exports / analytics", "Reporting and decision-support outputs.", "Use approved, filtered, correctly scoped data; exports honour role, scope, and mode."],
    ], widths=[1.7, 2.4, 2.4])

    # ---- 9. Roles ----
    h1(doc, "9. User Roles and Permissions")
    para(doc,
         "Sesigo uses two role concepts. The platform role (users.User.role) controls platform-level access. The "
         "project-scoped organisation role (projects.ProjectOrganization.role) controls how an organisation "
         "participates inside a specific project. The same organisation can play different roles in different projects.")
    h2(doc, "9.1 Platform roles (users.User.role)")
    table(doc, ["Role value", "Display", "Typical capability"], [
        ["admin", "Admin", "Full administration: users, projects, assignments, approvals, flags, training reset, settings, audit."],
        ["manager", "M&E Manager", "Oversight, targets, approvals, scheduled reports within assigned scope."],
        ["officer", "M&E Officer", "Operational editing, capture, flags, and reporting within allowed scope."],
        ["collector", "Data Collector", "Field-level data capture against assigned indicators/projects."],
        ["client", "Client", "Read-only access to approved dashboards, reports, and exports."],
    ], widths=[1.3, 1.6, 3.6])
    h2(doc, "9.2 Project-scoped organisation roles (projects.ProjectOrganization.role)")
    table(doc, ["Role value", "Meaning"], [
        ["lead", "Lead / overseer with project accountability."],
        ["coordinator", "Coordinates sub-grantees and may self-report."],
        ["sub_grantee", "Captures assigned indicators."],
        ["implementing_partner", "Direct service implementation and reporting."],
        ["data_reviewer", "Quality review and approval support."],
        ["funder", "Funding context and visibility."],
        ["other", "Any additional project-specific role."],
    ], widths=[2.0, 4.5])
    h2(doc, "9.3 Practical permission behaviour")
    bullets(doc, [
        "Admins can configure the system, manage users, create projects, assign organisations, review/approve data, and manage training reset activities.",
        "M&E managers and officers oversee or operate within their assigned project and organisation scope.",
        "Coordinators can manage sub-grantees and may also submit their own assigned indicators.",
        "Sub-grantees capture data only for indicators assigned to them in the project.",
        "Clients/funders normally access approved dashboards, reports, and exports without editing operational data.",
        "The backend enforces permissions even if the frontend hides buttons or screens; sensitive actions are recorded in UserActivity.",
    ])

    # ---- 10. Data capture / indicators ----
    h1(doc, "10. Data Capture, Indicators, and Disaggregation")
    para(doc,
         "Indicators are the centre of reporting. A project defines which indicators it tracks, sets targets (including "
         "per-organisation targets), defines disaggregation rules, and assigns those indicators to the correct "
         "coordinator or sub-grantee. Data capture follows the project indicator assignment rather than allowing "
         "unrestricted reporting.")
    h2(doc, "10.1 Indicator structure")
    para(doc,
         "The structure is: standard Indicator → ProjectIndicator (with targets + disaggregation rules) → "
         "ProjectIndicatorAssignment (organisation) → period data (Aggregate / Interaction+Response) → review "
         "status → dashboard output. The same indicator is reusable across projects while targets, assignments, and "
         "disaggregation differ per project.")
    h2(doc, "10.2 Indicator types and categories")
    table(doc, ["Aspect", "Supported values (from Indicator model)"], [
        ["Indicator types", "yes_no, number, percentage, text, select, multiselect, date, multi_int (multiple integers)."],
        ["Indicator categories", "hiv_prevention, ncd, mental_health, gbv, sti, trainings, media, events."],
        ["Disaggregation", "aggregate_disaggregation_config on the indicator plus ProjectIndicatorDisaggregationRule define the dimensions/values captured (e.g. age, sex, key population)."],
        ["Aliases", "IndicatorAlias maps alternate import/display names to a canonical indicator to prevent duplicates."],
    ], widths=[1.8, 4.7])
    h2(doc, "10.3 Cross-cutting indicators")
    para(doc,
         "Some health areas cut across target groups (for example, people living with HIV may also need NCD education). "
         "The portal avoids duplicate indicators with slightly different wording: it uses canonical indicator names and "
         "indicator aliases, and captures target group, health area, age, sex, and key population as disaggregation or "
         "classification fields where appropriate.")
    h2(doc, "10.4 Important programme areas and target groups")
    table(doc, ["Area", "Examples / notes"], [
        ["Health areas", "HIV prevention, NCD, mental health, STI, GBV, human rights, substance/alcohol abuse, sexual reproductive health."],
        ["Coordinator clusters", "Male engagement and faith-based, people living with HIV, NCD, adolescents and young people, testing and counselling, and key/vulnerable populations."],
        ["Age disaggregation", "Use standard age bands; adolescents and young people can be calculated from ages 10-24 where required."],
        ["Sex disaggregation", "Use sex fields to calculate male engagement and other sex-specific outputs."],
        ["PLHIV disaggregation", "Allow PLHIV views by age, sex, key population, and service area where captured."],
    ], widths=[1.8, 4.7])

    # ---- 11. Data quality, review, flagging (NEW) ----
    h1(doc, "11. Data Quality, Review, and Flagging")
    para(doc,
         "Sesigo treats data quality as a first-class workflow, not an afterthought. Submitted aggregate records move "
         "through an explicit review lifecycle, and any record can be flagged for follow-up.")
    h2(doc, "11.1 Aggregate review lifecycle (Aggregate.status)")
    table(doc, ["Status", "Meaning"], [
        ["draft", "Work in progress; not yet submitted for review."],
        ["pending", "Submitted and awaiting review."],
        ["reviewed", "Checked by a reviewer."],
        ["flagged", "A data-quality or follow-up issue has been raised."],
        ["approved", "Accepted as official data; eligible for dashboards/exports."],
        ["rejected", "Not accepted; must be corrected and resubmitted."],
    ], widths=[1.5, 5.0])
    h2(doc, "11.2 Flags (flags app)")
    para(doc,
         "A Flag can be attached to any record (by content_type + object_id) and is scoped to an organisation. Flags "
         "drive the Data Quality area of the portal.")
    table(doc, ["Attribute", "Values / behaviour"], [
        ["Flag type", "data_quality, follow_up, urgent, review, other."],
        ["Status", "open, in_progress, resolved, dismissed."],
        ["Priority", "low, medium, high, critical."],
        ["Assignment", "Can be assigned to a user; resolution captures resolved_by, resolved_at, and resolution_notes."],
        ["Discussion", "FlagComment records allow threaded comments on a flag."],
    ], widths=[1.6, 4.9])
    para(doc,
         "The frontend exposes this through the Flags and Data Quality pages, where managers and officers triage, "
         "assign, comment on, and resolve issues before data is approved.")

    # ---- 12. Reports/dashboards ----
    h1(doc, "12. Reports, Dashboards, Analytics, and Exports")
    para(doc,
         "Reports and dashboards are only as strong as the data pipeline behind them. Sesigo presents approved, scoped, "
         "and correctly filtered data. Users can filter by project, organisation, indicator, reporting period, age, sex, "
         "key population, and other approved disaggregation fields.")
    h2(doc, "12.1 Dashboard responsibilities")
    bullets(doc, [
        "Show project-level performance against targets (including coordinator targets).",
        "Show coordinator and sub-grantee contribution clearly.",
        "Allow filtering without mixing unrelated projects or organisations.",
        "Avoid double-counting where coordinator and sub-grantee reporting overlap.",
        "Keep training data out of live dashboards unless an authorised admin uses an include_training read-only view.",
    ])
    h2(doc, "12.2 Analysis features (analysis app)")
    table(doc, ["Feature", "Purpose"], [
        ["Report", "Saved or generated reports for reuse and sharing."],
        ["SavedQuery", "Reusable query definitions for quick repeat analysis."],
        ["ScheduledReport", "Report definitions that run on a schedule."],
        ["CoordinatorTarget", "Coordinator-level targets used in target-vs-achievement views."],
    ], widths=[1.8, 4.7])
    h2(doc, "12.3 Export responsibilities")
    bullets(doc, [
        "Exports respect the same permissions and filters as the on-screen dashboard.",
        "Exports identify project, organisation, indicator, reporting period, and approval status.",
        "Exports use standard labels so data can be reused for funder reporting and programme review.",
        "Generated exports are tracked as ExportJob records and recorded in the activity log.",
    ])

    # ---- 13. Imports/exports/workbooks (NEW) ----
    h1(doc, "13. Imports, Exports, and Report Workbooks")
    para(doc,
         "Sesigo supports bulk data movement so partners are not limited to single-record entry. The uploads app tracks "
         "the underlying jobs and files, and the frontend provides workbook and batch tools.")
    table(doc, ["Capability", "What it does"], [
        ["Upload", "Stores uploaded files (e.g. narrative reports, supporting documents) linked to scope."],
        ["ImportJob", "Tracks a bulk data import, including validation against project/organisation/indicator scope and aliases."],
        ["ExportJob", "Tracks a generated export file for download and audit."],
        ["Report workbooks (frontend)", "Workbook-style data entry/import for structured reporting (report-workbooks page)."],
        ["Batch record (frontend)", "Bulk capture of multiple records in one operation (batch-record page)."],
        ["Check-in (frontend)", "Lightweight field check-in flow for collectors (checkin route)."],
    ], widths=[2.0, 4.5])
    para(doc,
         "Imports use the same boundary rules as manual capture: organisation scope, project scope, indicator "
         "assignment, and live/training mode are all enforced before records are created.")

    # ---- 14. Messaging/social (NEW) ----
    h1(doc, "14. Messaging, Announcements, Notifications, and Social Tracking")
    para(doc,
         "Sesigo includes lightweight collaboration features so coordination does not have to happen outside the "
         "system.")
    table(doc, ["Feature (model)", "Purpose"], [
        ["Message (messaging)", "Internal direct messaging between users."],
        ["Announcement (messaging)", "System-wide or scoped announcements (Announcements page)."],
        ["Notification (messaging)", "System notifications for events such as approvals, flags, and assignments."],
        ["SocialPost (social)", "Tracks social-media posts against an indicator, so media/communication outputs count toward results."],
    ], widths=[2.0, 4.5])
    para(doc,
         "These features are surfaced in the frontend Messages, Announcements, and Social pages, with a global Search "
         "page and a System Status page for operational visibility.")

    # ---- 15. Live vs training ----
    h1(doc, "15. Sesigo Live System vs Sesigo Training Mode")
    para(doc,
         "Sesigo has two operating modes. The Sesigo Live System is used for official production data. Sesigo Training "
         "Mode is used for practice, onboarding, testing, and demonstrations. The modes run on the same codebase, but "
         "the data boundary is enforced at the backend.")
    diagram_flow(doc, [
        "Login / mode selection",
        "Live routes (/dashboard, /reports, /aggregates)   |   Training routes (/training/...)",
        "API request — mode-aware signal (training_only=true for /training)",
        "Backend mode guard — apply_training_filter(...) + assert_project_write_allowed(...)",
        "Live data (is_training=false)   |   Training data (is_training=true)",
        "Rule: training data must never mix with live data",
    ])
    caption(doc, "Figure 4: Live and training separation.")
    table(doc, ["Concern", "Sesigo Live System", "Sesigo Training Mode"], [
        ["Route pattern", "/dashboard, /reports, /aggregates, /analysis, /projects", "/training/dashboard, /training/reports, /training/aggregates, /training/analysis, /training/projects"],
        ["Request signal", "No training flag by default.", "training_only=true or mode=training request logic."],
        ["Project scope", "is_training = false.", "is_training = true."],
        ["Read behaviour", "Excludes training-linked records by default.", "Returns training-linked records only."],
        ["Write behaviour", "Cannot write into training projects.", "Cannot write into live projects."],
        ["Shared lists", "Hide orgs/indicators/clients linked only to training projects.", "Show training-linked orgs/indicators/clients."],
        ["Admin override", "Read-only include_training=true may expose both.", "Same read-all override may apply, but writes stay blocked across boundaries."],
    ], widths=[1.4, 2.4, 2.7])
    h2(doc, "15.1 Why backend enforcement is important")
    para(doc,
         "The frontend route is helpful but not sufficient. A user, browser error, or faulty request can send the wrong "
         "URL or query string. The backend therefore always checks the target project and rejects cross-mode writes via "
         "assert_project_write_allowed(...). Reads are filtered with apply_training_filter(...), "
         "apply_training_filter_to_projects(...), and apply_training_filter_via_projects(...). Server-side write guards "
         "are a non-negotiable design rule, and a cleanup_training_data management command can reset training records.")

    # ---- 16. DB model ----
    h1(doc, "16. Simplified Database Model")
    para(doc,
         "The database model is centred on the Project. The Project provides the reporting boundary. ProjectOrganization "
         "defines who participates and in what role. ProjectIndicator defines what is measured (with targets and "
         "disaggregation rules). ProjectIndicatorAssignment controls who reports which indicator. Aggregate or "
         "Interaction/Response records carry the data that feeds reports.")
    table(doc, ["Entity", "Purpose"], [
        ["User / UserActivity", "User account, platform role, organisation link; activity/audit trail."],
        ["Organization", "BONASO, coordinators, sub-grantees, and implementing partners (with parent links)."],
        ["ClientOrganization", "Funding source linked to projects."],
        ["Project", "Funded project, status, and live/training (is_training) flag; hierarchy_overrides fallback."],
        ["ProjectOrganization", "Organisation role inside a specific project (coordinator/sub-grantee/implementer flags)."],
        ["ProjectOrganizationHierarchy", "Explicit parent-child organisation links within a project."],
        ["Indicator / IndicatorAlias", "Standard indicator definition plus alternate names resolving to it."],
        ["Assessment / AssessmentIndicator", "Group of indicators collected together."],
        ["ProjectIndicator (+ target/rule models)", "Indicator scoped to a project with targets, per-organisation targets, and disaggregation rules."],
        ["ProjectIndicatorAssignment", "Controls which organisation can report which project indicator."],
        ["Aggregate", "Summarised indicator results by period, project, organisation, and status."],
        ["Respondent / Interaction / Response", "Individual-level beneficiaries, service interactions, and indicator responses."],
        ["Profile / ProfileField", "Extended/custom profile fields for respondents and participants."],
        ["Event / Participant / EventPhase", "Activity/event participation data and phases."],
        ["SocialPost", "Social-media post tracked against an indicator."],
        ["Flag / FlagComment", "Data-quality/follow-up flags on any record, with comments."],
        ["NarrativeReport", "Narrative reports/documents linked to project and organisation."],
        ["Report / SavedQuery / ScheduledReport / CoordinatorTarget", "Analysis and reporting artefacts."],
        ["Upload / ImportJob / ExportJob", "Files and tracked import/export jobs."],
        ["Message / Announcement / Notification", "Internal communication and notifications."],
    ], widths=[2.4, 4.1])

    # ---- 17. Usage by group ----
    h1(doc, "17. How the System Is Used by Each User Group")
    h2(doc, "17.1 System administrator")
    bullets(doc, [
        "Create and manage users, roles, and permissions.",
        "Create clients/funders and projects; set live or training.",
        "Assign organisations to projects and define hierarchy.",
        "Configure indicators, aliases, assessments, targets, and organisation assignments.",
        "Monitor activity logs, flags, backups, training reset/cleanup, and security settings.",
    ])
    h2(doc, "17.2 M&E manager / officer")
    bullets(doc, [
        "Review project performance across coordinators and sub-grantees.",
        "Monitor reporting completeness and quality; triage and resolve flags.",
        "Approve, query, or reject submitted data.",
        "Set coordinator targets and use dashboards to identify gaps and trends.",
        "Schedule and export approved data for donor and management reporting.",
    ])
    h2(doc, "17.3 Coordinator organisation")
    bullets(doc, [
        "View assigned project and sub-grantee structure.",
        "Submit coordinator-level data where the coordinator implements directly.",
        "Monitor sub-grantee submissions and follow up missing reports and flags.",
        "Use dashboards to see own contribution and sub-grantee performance.",
    ])
    h2(doc, "17.4 Sub-grantee organisation")
    bullets(doc, [
        "Log in to the assigned project scope.",
        "Capture data against assigned indicators (single, batch, or workbook).",
        "Upload supporting narrative reports where required.",
        "Review submission status and correct rejected or flagged records.",
    ])
    h2(doc, "17.5 Trainer or trainee")
    bullets(doc, [
        "Use Sesigo Training Mode (/training routes) for safe practice.",
        "Enter demo records without affecting official live data.",
        "Practice dashboards, reports, and exports using training records only.",
        "Training data can be reset/purged (cleanup_training_data) without affecting production data.",
    ])

    # ---- 18. Admin/security/QA ----
    h1(doc, "18. Administration, Security, and Quality Assurance")
    h2(doc, "18.1 Administration responsibilities")
    bullets(doc, [
        "Keep user accounts and roles accurate.",
        "Keep project and organisation structures aligned with signed implementation arrangements.",
        "Maintain indicator catalogues and aliases to prevent duplicate wording.",
        "Back up the database and media files regularly.",
        "Document deployments, configuration changes, and migrations.",
    ])
    h2(doc, "18.2 Security responsibilities")
    bullets(doc, [
        "Use secure production settings for HTTPS, cookies, CSRF, and secret keys.",
        "Authenticate with JWT (SimpleJWT) and enforce backend permissions for every read and write.",
        "Keep training/live isolation server-side, not only through the UI.",
        "Apply least-privilege access for clients, coordinators, sub-grantees, and collectors.",
        "Protect respondent-level data and avoid unnecessary exposure in exports.",
        "Use UserActivity logs to trace who created, edited, approved, exported, or imported records.",
    ])
    h2(doc, "18.3 Quality assurance responsibilities")
    bullets(doc, [
        "Test dashboard totals against known datasets.",
        "Test organisation scope and project-specific indicator assignments.",
        "Test training-mode isolation for reads, writes, dashboards, exports, respondents, and events (tests_training_separation).",
        "Test import templates, alias resolution, and validation messages before rollout.",
        "Use the flag lifecycle and audit logs to trace and close data-quality issues.",
    ])

    # ---- 19. Dev checklist ----
    h1(doc, "19. Implementation Checklist for Developers")
    table(doc, ["Area", "Checklist item"], [
        ["Project model", "Every core reporting record must link to a project or be filterable through a project-linked entity."],
        ["Training/live isolation", "Every read endpoint applies a training filter; every write endpoint calls assert_project_write_allowed."],
        ["Organisation scope", "Non-admin users only see organisations within their allowed project tree (organizations/access.py)."],
        ["Indicator assignment", "Capture forms and imports check that the organisation has the indicator assignment."],
        ["Dashboards", "Dashboard queries avoid mixing training/live data and avoid double-counting."],
        ["Exports", "Exports apply the same filters as the UI/API and create an ExportJob record."],
        ["Imports", "ImportJob validation enforces scope, assignment, mode, and alias resolution."],
        ["Uploads", "Uploaded reports link to project, organisation, user, and mode."],
        ["Flags", "Data-quality issues use the flag lifecycle (open → resolved/dismissed) before approval."],
        ["Audit logs", "Critical actions record user, timestamp, action, and affected record (UserActivity)."],
        ["Testing", "Cover permission boundaries, mode separation, alias import, and dashboard totals."],
        ["Documentation", "Re-run this generator and update system maps/ERDs when schema or workflows change."],
    ], widths=[1.8, 4.7])

    # ---- 20. Future improvements ----
    h1(doc, "20. Future Improvements")
    para(doc,
         "Several items previously listed as future work are now implemented (a user-activity audit trail, indicator "
         "aliasing, and assignment-driven reporting). Remaining opportunities include:")
    bullets(doc, [
        "Add a physically separate training database/deployment for stronger infrastructure isolation (a training/compose stack already exists as scaffolding).",
        "Extend the audit trail to capture old/new field values (diffs) on update for richer change history.",
        "Add configurable reporting periods, quarterly targets, and richer target-versus-achievement views.",
        "Add automated reset/expiry rules for training data (e.g. scheduled purge after N days).",
        "Add more user-facing help text and tooltips for coordinators and sub-grantees.",
        "Add automated QA checks for completeness, outliers, duplicate submissions, and invalid disaggregation.",
        "Continue consolidating duplicate indicators through aliases and canonical labels.",
    ])

    # ---- 21. Conclusion ----
    h1(doc, "21. Architectural Conclusion")
    para(doc,
         "Sesigo Data Portal should be understood as a project-scoped reporting platform. Its strength comes from "
         "explicit modelling of the funder, project, BONASO oversight, coordinator organisations, sub-grantees, "
         "indicators, assignments, data-quality flags, validation, approval, communication, and reporting outputs.")
    para(doc,
         "The long-term technical direction keeps three principles at the centre: project scope must be explicit, "
         "indicator assignment must control reporting, and training data must remain isolated from live official "
         "reporting data.")
    para(doc,
         "When new features are added, developers should always ask: Which project does this belong to? Which "
         "organisation is allowed to see or edit it? Which indicator assignment permits the record? Is this live data or "
         "training data? Answering those consistently keeps the Sesigo Data Portal scalable, auditable, and safe for "
         "BONASO programme reporting.")

    doc.save(str(OUT))
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
