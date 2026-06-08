"""
Share-ready overview/mapping Word document for colleagues: presents the
integrated community assessment as a reviewable Question -> Indicator ->
Organisation mapping (plus per-organisation sets and the programme-indicator
appendix). Reuses the question bank from build_assessment_form.py.
"""
import os
import sys
from pathlib import Path

BACKEND_ROOT = Path(__file__).resolve().parents[1]
if str(BACKEND_ROOT) not in sys.path:
    sys.path.insert(0, str(BACKEND_ROOT))
os.environ.setdefault("DJANGO_SETTINGS_MODULE", "core.settings")
import django  # noqa: E402

django.setup()

from indicators.models import Indicator  # noqa: E402
from docx import Document  # noqa: E402
from docx.shared import Pt, Inches, RGBColor  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.enum.section import WD_ORIENT  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.oxml import OxmlElement  # noqa: E402

import build_assessment_form as B  # noqa: E402

NAVY = RGBColor(0x1F, 0x38, 0x64)
GREY = RGBColor(0x55, 0x55, 0x55)
HDRBG = "1F3864"
MODBG = "D9E1F2"


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hexcolor)
    tcPr.append(sh)


def setcell(cell, text, *, bold=False, size=8.5, color=None, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color


def main():
    org_map = B.load_org_map()
    code = {i: c for i, c in Indicator.objects.values_list("id", "code")}
    name = {i: n for i, n in Indicator.objects.values_list("id", "name")}

    doc = Document()
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = sec.page_height, sec.page_width
    for m in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(sec, m, Inches(0.6))
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    def heading(text, size=14, before=10, after=4, color=NAVY):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after = Pt(after)
        r = p.add_run(text); r.bold = True; r.font.size = Pt(size); r.font.color.rgb = color
        return p

    # Title
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("NAHPA Social Contracting 2026/27"); r.bold = True; r.font.size = Pt(18); r.font.color.rgb = NAVY
    t2 = doc.add_paragraph(); t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t2.add_run("Integrated Community Assessment — Question / Indicator / Organisation Mapping")
    r.bold = True; r.font.size = Pt(13)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("For review. A single individual-assessment questionnaire whose answers feed monitoring "
                  "indicators; each indicator belongs to one or more coordinating organisations, so every "
                  "organisation collects only the questions tied to its indicators.")
    r.italic = True; r.font.size = Pt(9.5); r.font.color.rgb = GREY

    def rtype_label(it):
        k = it["kind"]
        if k == "check":
            return "Tick (done)"
        if k == "yesno":
            return "Yes / No"
        if k == "num":
            return "Number"
        if k == "select":
            return "Select: " + " / ".join(it["options"])
        return ""

    items = [(title, it) for title, it in B.all_indicator_items()]
    # Overview counts
    n_ids = len({i for _, it in items for i in it["ids"]})
    heading("Overview", 12)
    for line in [
        f"• {len(items)} screening/service items across {len(B.IND_SECTIONS)} sections, "
        f"plus Session-info and Demographics (BONELA-style layout).",
        f"• Mapped to {n_ids} individual-level indicators (shared indicators collapse into one item — "
        "e.g. all population-specific condom indicators → one ‘Male condoms — quantity’ entry).",
        f"• {len(B.PROGRAMME_IDS)} programme/activity indicators are NOT individual items (Appendix A).",
        "• Response types: Tick (service done), Yes/No, Number, Select-one.",
        "• Disaggregation captured once per client: sex, age band, key population, district/locality.",
    ]:
        b = doc.add_paragraph(line); b.paragraph_format.space_after = Pt(1); b.runs[0].font.size = Pt(9.5)

    # ---- Master mapping table ----
    heading("Item – Indicator – Organisation mapping", 12)
    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = "Table Grid"
    widths = [Inches(0.4), Inches(3.6), Inches(1.6), Inches(3.3), Inches(1.55)]
    hdrs = ["#", "Screening / service item", "Response", "Indicator (code – name)", "Organisation(s)"]
    for c, h, w in zip(tbl.rows[0].cells, hdrs, widths):
        setcell(c, h, bold=True, size=9, color=RGBColor(0xFF, 0xFF, 0xFF))
        shade(c, HDRBG)
        c.width = w

    qnum = 0
    seen_section = None
    qmap = []
    for title, it in items:
        if title != seen_section:
            seen_section = title
            row = tbl.add_row().cells
            a = row[0]
            a.merge(row[1]); a.merge(row[2]); a.merge(row[3]); a.merge(row[4])
            setcell(a, title, bold=True, size=9.5, color=NAVY)
            shade(a, MODBG)
        qnum += 1
        ids = it["ids"]
        ogs = B.orgs_for(ids, org_map)
        qmap.append((qnum, ids, ogs))
        cells = tbl.add_row().cells
        for c, w in zip(cells, widths):
            c.width = w
        setcell(cells[0], str(qnum), bold=True, size=8.5)
        setcell(cells[1], it["label"].lstrip("↳ "), size=8.5)
        setcell(cells[2], rtype_label(it), size=8)
        refs = "\n".join(f"{code.get(i,'?')} – {name.get(i,'?')}" for i in ids)
        setcell(cells[3], refs, size=7.5)
        setcell(cells[4], ", ".join(ogs), size=8)

    # ---- Per-organisation sets ----
    doc.add_page_break()
    heading("Each organisation's question set", 12)
    ot = doc.add_table(rows=1, cols=4); ot.style = "Table Grid"
    ow = [Inches(2.0), Inches(3.6), Inches(0.8), Inches(4.0)]
    for c, h, w in zip(ot.rows[0].cells, ["Organisation", "Programme focus", "Q count", "Question numbers"], ow):
        setcell(c, h, bold=True, size=9, color=RGBColor(0xFF, 0xFF, 0xFF)); shade(c, HDRBG); c.width = w
    for org in B.COORD_ORDER:
        qs = [q for q, ids, ogs in qmap if org in ogs]
        cells = ot.add_row().cells
        for c, w in zip(cells, ow):
            c.width = w
        setcell(cells[0], org, bold=True, size=9)
        setcell(cells[1], B.COORD_FOCUS[org], size=8.5)
        setcell(cells[2], str(len(qs)), size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        setcell(cells[3], ", ".join(f"Q{q}" for q in qs), size=8)

    # ---- Appendix ----
    doc.add_page_break()
    heading("Appendix A — Programme / activity indicators (not collected per individual)", 12)
    doc.add_paragraph("Captured through activity / monthly reports, not the individual questionnaire.").runs[0].italic = True
    at = doc.add_table(rows=1, cols=3); at.style = "Table Grid"
    aw = [Inches(3.6), Inches(5.4), Inches(1.4)]
    for c, h, w in zip(at.rows[0].cells, ["Indicator code", "Indicator", "Organisation(s)"], aw):
        setcell(c, h, bold=True, size=9, color=RGBColor(0xFF, 0xFF, 0xFF)); shade(c, HDRBG); c.width = w
    for i in B.PROGRAMME_IDS:
        cells = at.add_row().cells
        for c, w in zip(cells, aw):
            c.width = w
        setcell(cells[0], code.get(i, "?"), size=8)
        setcell(cells[1], name.get(i, "?"), size=8)
        setcell(cells[2], ", ".join(B.orgs_for([i], org_map)), size=8)

    out = BACKEND_ROOT.parent / "frontend/docs/NAHPA_SC_2026-27_Assessment_Mapping_for_Review.docx"
    doc.save(str(out))
    print("Saved:", out)
    print(f"questions: {qnum} | programme indicators: {len(B.PROGRAMME_IDS)}")


if __name__ == "__main__":
    main()
