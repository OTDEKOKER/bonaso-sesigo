"""
Build the NAHPA Social Contracting 2026/27 Community Health Worker (CHW)
Individual Screening & Service Record as a Word document.

Modelled on the BONELA Individual Outreach Form (see
indicators/management/commands/seed_bonela_outreach_form.py): completed by the
CHW during a community screening / service session — Session info, Demographics,
then tick-list sections for Information provided, Screening conducted, Referrals
made, Linked for care, and Commodities distributed. Every indicator-bearing item
is linked to its indicator(s); indicator -> coordinator organisation membership
is read live from project 3 so each organisation's set is accurate.

Run (in venv):  python scripts/build_assessment_form.py
"""
import os
import sys
from pathlib import Path

BACKEND_ROOT = Path(__file__).resolve().parents[1]
if str(BACKEND_ROOT) not in sys.path:
    sys.path.insert(0, str(BACKEND_ROOT))
os.environ.setdefault("DJANGO_SETTINGS_MODULE", "core.settings")
import django  # noqa: E402

django.setup()

from indicators.models import Indicator  # noqa: E402
from projects.models import ProjectIndicatorOrganizationTarget as OT  # noqa: E402

from docx import Document  # noqa: E402
from docx.shared import Pt, Inches, RGBColor  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402

COORD = {166: "MBGE", 5: "MAKGABANENG", 112: "BONEPWA+", 109: "HPP", 1: "TEBELOPELE", 159: "MOPIPI"}
COORD_ORDER = ["MBGE", "MAKGABANENG", "BONEPWA+", "HPP", "TEBELOPELE", "MOPIPI"]
COORD_FOCUS = {
    "MBGE": "Men & Boys / Male engagement",
    "MAKGABANENG": "Non-communicable diseases (NCD), cancer & mental health",
    "BONEPWA+": "People living with HIV (PLHIV)",
    "HPP": "Key & vulnerable populations (KVP)",
    "TEBELOPELE": "HIV testing services",
    "MOPIPI": "Adolescents & young people (AYP)",
}

# ---- context sections (no indicators): (title, note, [(label, widget)]) ----
CONTEXT_SECTIONS = [
    ("Section A — Session Information",
     "Completed by the Community Health Worker for each community screening / service session.",
     [
         ("Implementing organisation (coordinator)",
          "MBGE ☐   MAKGABANENG ☐   BONEPWA+ ☐   HPP ☐   TEBELOPELE ☐   MOPIPI ☐"),
         ("Sub-grantee / service point", "________________________________"),
         ("District", "____________________     Village / locality: ____________________"),
         ("Session / service location",
          "Household ☐   Outreach site ☐   Health post ☐   Hotspot ☐   Facility ☐   Other: __________"),
         ("Date of session", "____ / ____ / 20____"),
         ("Community Health Worker / Mobiliser (name & ID)", "________________________________"),
     ]),
    ("Section B — Client Demographics & Disaggregation",
     "Recorded once per client; disaggregates every indicator below.",
     [
         ("Unique client ID / code", "________________________________"),
         ("Sex", "Male ☐      Female ☐      Intersex ☐"),
         ("Age (years)", "______     Band:  <15 ☐   15–24 ☐   25–49 ☐   50+ ☐"),
         ("Key / vulnerable population (tick all that apply)",
          "General population ☐   FSW ☐   MSM ☐   PWD (disability) ☐   LGBTQ+ ☐   PWID ☐   PLHIV ☐   AYP (10–24) ☐"),
         ("Relationship status", "Single ☐   Married ☐   Cohabiting ☐   Widowed ☐   Divorced ☐"),
         ("Citizenship", "Citizen ☐      Non-citizen ☐"),
         ("Pregnancy status (where applicable)", "Yes ☐      No ☐      N/A ☐"),
     ]),
]

# item kinds: "check" (tick one box -> indicator), "num", "yesno",
#             "select" (choose one of options)
def chk(label, ids):
    return {"kind": "check", "label": label, "ids": ids}


def num(label, ids):
    return {"kind": "num", "label": label, "ids": ids}


def yn(label, ids):
    return {"kind": "yesno", "label": label, "ids": ids}


def sel(label, ids, options):
    return {"kind": "select", "label": label, "ids": ids, "options": options}


# ---- indicator sections (BONELA-style grouping): (title, note, [items]) ----
IND_SECTIONS = [
    ("Section C — Information & Education Provided",
     "Tick each message / education topic provided to the client during the session.",
     [
         chk("HIV prevention messages", [558]),
         chk("Stigma-reduction messages", [330]),
         chk("NCD prevention information", [372]),
         chk("NCD prevention messages via social media", [373]),
         chk("Basic human rights & HIV", [358]),
         chk("Self-breast examination education", [374]),
         chk("Prostate cancer education", [375]),
         chk("Cervical cancer education", [376]),
     ]),
    ("Section D — Screening Conducted",
     "Tick each screening performed; record results where shown.",
     [
         chk("Gender-based violence (GBV) screening", [343]),
         chk("Eligible for GBV services (based on screening)", [462]),
         chk("STI screening", [351]),
         chk("Tobacco-use screening (NCD risk)", [420]),
         chk("Alcohol-use screening (NCD risk)", [425]),
         chk("NCD risk factors — blood glucose / blood pressure / BMI", [486]),
         chk("Breast cancer screening", [380]),
         chk("Mental health screening", [385]),
         chk("HIV test conducted", [451]),
         sel("↳ HIV test result", [452], ["Negative", "Positive", "Declined", "Not tested"]),
         sel("↳ TB screening outcome (PLHIV)", [525],
             ["Not screened", "Screened–negative", "Positive & on treatment"]),
     ]),
    ("Section E — Referrals Made",
     "Tick each referral made for the client.",
     [
         chk("HIV testing", [445]),
         chk("PrEP (pre-exposure prophylaxis)", [337]),
         chk("PEP (post-exposure prophylaxis)", [446]),
         chk("GBV — clinical services", [464]),
         chk("GBV — justice services", [466]),
         chk("STI services", [550]),
         chk("STI services (screened positive)", [468]),
         chk("Further screening services", [388]),
         chk("Diabetes treatment / management", [384]),
         chk("Mental health management / treatment", [389]),
         chk("Further health services (NCD risk suspected)", [383]),
         chk("Legal aid services", [473]),
         chk("Justice services", [474]),
     ]),
    ("Section F — Linked for Care / Services",
     "Tick each service the client was linked to / received.",
     [
         chk("HIV care & treatment (if HIV-positive)", [453]),
         chk("PLHIV — treatment, care & support", [484]),
         chk("ART reinitiation (treatment had been interrupted)", [483]),
         chk("Treatment literacy (PLHIV)", [481]),
         chk("Legal services (PLHIV)", [485]),
         chk("Psychosocial support (GBV)", [350]),
         chk("Mental health counselling", [386]),
         chk("Family planning services (AYP)", [342]),
         chk("Community physical-activity club", [390]),
         chk("Tobacco cessation programme", [487]),
         chk("Alcohol-abuse programme", [382]),
         chk("STI referral completed (attended service)", [470]),
         chk("Person with disability accessed services", [628]),
         chk("Key population member accessed services", [629]),
     ]),
    ("Section G — Commodities Distributed",
     "Record quantities given to the client.",
     [
         num("Male condoms — quantity", [471, 355, 447, 627]),
         num("Braille-labelled condoms — quantity", [448]),
         num("Lubricant sachets — quantity", [357]),
         yn("Repeat condom collection (client has collected condoms before)", [356]),
         yn("Distributed through a non-traditional outlet (bar, hotspot, peer)", [450]),
     ]),
    ("Section H — Human Rights & Health Events",
     "Tick where applicable.",
     [
         yn("Rights violation reported and redress sought", [538]),
         chk("Event attended: SADC Healthy Lifestyle commemoration", [392]),
         chk("Event attended: World Diabetes Day", [393]),
         chk("Event attended: Breast Cancer Month", [394]),
         chk("Event attended: Cervical Cancer Awareness", [395]),
         chk("Event attended: Mental Health Day", [396]),
         chk("Event attended: World Cancer Day", [397]),
         chk("Event attended: Mo-Vember", [398]),
     ]),
]

# programme / activity-level indicators -> appendix (NOT individual items)
PROGRAMME_IDS = [365, 367, 363, 364, 368, 366, 370, 318, 320, 631, 319, 630,
                 391, 521, 442, 359, 444, 362, 438, 626, 387]

NAVY = RGBColor(0x1F, 0x38, 0x64)
GREY = RGBColor(0x55, 0x55, 0x55)


def load_org_map():
    m = {}
    for ot in OT.objects.filter(project_indicator__project_id=3, organization_id__in=COORD):
        m.setdefault(ot.project_indicator.indicator_id, set()).add(COORD[ot.organization_id])
    return m


def orgs_for(ids, m):
    s = set()
    for i in ids:
        s |= m.get(i, set())
    return [c for c in COORD_ORDER if c in s]


def all_indicator_items():
    """flatten indicator-bearing items across sections, in order."""
    for title, note, items in IND_SECTIONS:
        for it in items:
            yield title, it


def main():
    org_map = load_org_map()
    code = {i: c for i, c in Indicator.objects.values_list("id", "code")}
    name = {i: n for i, n in Indicator.objects.values_list("id", "name")}

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    def heading(text, size=14, color=NAVY, before=10, after=3):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.space_after = Pt(after)
        r = p.add_run(text); r.bold = True; r.font.size = Pt(size); r.font.color.rgb = color
        return p

    def ref(ids):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_after = Pt(5)
        refs = "; ".join(f"{code.get(i,'?')}" for i in ids)
        og = ", ".join(orgs_for(ids, org_map)) or "(no coordinator)"
        r = p.add_run(f"→ Indicator: {refs}   |   Org(s): {og}")
        r.italic = True; r.font.size = Pt(7.5); r.font.color.rgb = GREY

    # Title
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("NAHPA Social Contracting 2026/27"); r.bold = True; r.font.size = Pt(19); r.font.color.rgb = NAVY
    t2 = doc.add_paragraph(); t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t2.add_run("Individual Assessment Questionnaire — Community Screening & Service Record")
    r.bold = True; r.font.size = Pt(13)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Completed by the Community Health Worker for each individual during community "
                  "screening and service delivery. The grey line under each item shows the "
                  "indicator code(s) it populates and the organisation(s) that report it.")
    r.italic = True; r.font.size = Pt(9.5); r.font.color.rgb = GREY

    heading("How responses populate the indicators", 11, NAVY, before=8, after=2)
    for line in [
        "• Each item is linked to one indicator (some shared across organisations). A tick or “Yes” = one person counted; a quantity (condoms, lubricants) is added up.",
        "• Every count is automatically disaggregated by the client’s sex, age band and key population captured once in Section B — giving the total / male / female / age / key-population breakdown the aggregate report needs.",
        "• Roll-up rule: tick & Yes/No items → COUNT of people; quantity items → SUM; result selects (HIV result, TB outcome) → COUNT within the chosen category.",
        "• Summed across all clients for the reporting period, these answers become each indicator’s reported value (actual) — i.e. the organisation’s aggregate report, measured against its targets.",
    ]:
        b = doc.add_paragraph(line); b.paragraph_format.space_after = Pt(1); b.runs[0].font.size = Pt(9)

    # Context sections (tables)
    for title, note, fields in CONTEXT_SECTIONS:
        heading(title, 14)
        doc.add_paragraph(note).runs[0].italic = True
        tbl = doc.add_table(rows=0, cols=2); tbl.style = "Light Grid Accent 1"
        for k, v in fields:
            cells = tbl.add_row().cells
            cells[0].width = Inches(2.4); cells[1].width = Inches(4.3)
            rr = cells[0].paragraphs[0].add_run(k); rr.bold = True; rr.font.size = Pt(9.5)
            cells[1].paragraphs[0].add_run(v).font.size = Pt(9.5)

    # Indicator sections
    for title, note, items in IND_SECTIONS:
        heading(title, 13, NAVY, before=12)
        doc.add_paragraph(note).runs[0].italic = True
        for it in items:
            line = doc.add_paragraph()
            line.paragraph_format.space_after = Pt(0)
            kind = it["kind"]
            if kind == "check":
                run = line.add_run("☐  " + it["label"]); run.font.size = Pt(10.5)
            elif kind == "yesno":
                line.add_run(it["label"] + "   ").font.size = Pt(10.5)
                line.add_run("Yes ☐   No ☐").font.size = Pt(10)
            elif kind == "num":
                line.add_run(it["label"] + "   ").font.size = Pt(10.5)
                line.add_run("|__________|").font.size = Pt(10)
            elif kind == "select":
                line.paragraph_format.left_indent = Inches(0.25)
                line.add_run(it["label"] + ":   ").font.size = Pt(10)
                line.add_run("   ".join(o + " ☐" for o in it["options"])).font.size = Pt(9.5)
            ref(it["ids"])

    # Per-organisation sets
    doc.add_page_break()
    heading("Each Organisation's Item Set", 14)
    doc.add_paragraph("The screening/service items above, grouped by the organisation that reports them. "
                      "A CHW working for an organisation completes that organisation's items.").runs[0].italic = True
    # build org -> list of (section short, label)
    for org in COORD_ORDER:
        items = [(title, it) for title, it in all_indicator_items() if org in orgs_for(it["ids"], org_map)]
        heading(f"{org}  —  {COORD_FOCUS[org]}   ({len(items)} items)", 12, NAVY, before=10, after=1)
        for title, it in items:
            sec = title.split("—")[1].strip()
            b = doc.add_paragraph(f"• [{sec}]  {it['label'].lstrip('↳ ')}")
            b.paragraph_format.space_after = Pt(0); b.runs[0].font.size = Pt(9)

    # Appendix
    doc.add_page_break()
    heading("Appendix A — Programme / Activity Indicators (not collected per individual)", 13)
    doc.add_paragraph("Organisational activities (staff training, mentorship, campaigns, support groups, "
                      "media, community-led monitoring, reporting) captured via activity / monthly reports — "
                      "not this individual record.").runs[0].italic = True
    at = doc.add_table(rows=1, cols=3); at.style = "Light Grid Accent 1"
    for c, h in zip(at.rows[0].cells, ["Indicator code", "Indicator", "Organisation(s)"]):
        rr = c.paragraphs[0].add_run(h); rr.bold = True; rr.font.size = Pt(9)
    for i in PROGRAMME_IDS:
        cells = at.add_row().cells
        cells[0].paragraphs[0].add_run(code.get(i, "?")).font.size = Pt(8)
        cells[1].paragraphs[0].add_run(name.get(i, "?")).font.size = Pt(8)
        cells[2].paragraphs[0].add_run(", ".join(orgs_for([i], org_map))).font.size = Pt(8)

    out = BACKEND_ROOT.parent / "frontend/docs/NAHPA_SC_2026-27_Individual_Assessment_Questionnaire.docx"
    doc.save(str(out))
    n_items = sum(len(items) for _, _, items in IND_SECTIONS)
    n_ids = len({i for _, it in all_indicator_items() for i in it["ids"]})
    print("Saved:", out)
    print(f"indicator items: {n_items} | individual indicators covered: {n_ids} | programme (appendix): {len(PROGRAMME_IDS)}")
    for org in COORD_ORDER:
        n = sum(1 for _, it in all_indicator_items() if org in orgs_for(it["ids"], org_map))
        print(f"  {org:12} {n} items")


if __name__ == "__main__":
    main()
